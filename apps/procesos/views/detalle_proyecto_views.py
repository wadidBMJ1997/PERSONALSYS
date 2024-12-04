from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from datetime import datetime
from openpyxl import load_workbook
from openpyxl.utils import range_boundaries, get_column_letter
from openpyxl.styles import Alignment
from docx.oxml.ns import qn
from django.http import HttpResponse
from django.views import View
from ..models.proyecto_models import DetalleProyecto, Proyecto  # Importa el modelo
from apps.maestros.models.persona_models import Persona, Cliente
import fitz
import io
import os
import zipfile
from io import BytesIO
from django.conf import settings

from django.shortcuts import get_object_or_404


class DetalleProyectoWordView(View):

    #Metodos de Cabin request

    def set_cell_border(self, cell):
        """
        Configura un borde alrededor de una celda de la tabla.
        """
        tc = cell._element.tcPr
        for border_name in ['top', 'start', 'bottom', 'end']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Tamaño del borde
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Color del borde (negro)
            tc.append(border)

    def set_font(self, cell, font_size=9.5, bold=False):
        """
        Establece la fuente de una celda.
        """
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
        run.font.size = Pt(font_size)
        run.bold = bold

    def add_centered_heading(self, doc, text, font_size):
        """
        Añade un encabezado centrado al documento.
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Método para generar todos los documentos 
    def get(self, request):

        #Carpeta contendeor general 
        
        # hay que validar cuando tenga mas de un código del  proyecto
        proyecto = Proyecto.objects.first()
        folder_general = proyecto.codigo_proyecto
        # Crear un objeto BytesIO para almacenar el archivo ZIP en memoria
        zip_buffer = BytesIO()

        # Crear el archivo ZIP
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            # --- Generar el archivo Excel ---
            file_path = os.path.join(settings.BASE_DIR, 'USS Technicians - Spreadsheet for Jobs.xlsx')

            # Cargar el archivo de Excel existente
            wb = load_workbook(file_path)
            ws = wb.active  # Hoja activa
            
            specific_data = {
                "Project": proyecto.codigo_proyecto,
                "Vessel": proyecto.vessel_name,
                "Start Date": proyecto.fecha_inicio_proyecto,
                "End Date": proyecto.fecha_final_proyecto,
                "Start Location": proyecto.ubicacion_inicial,
                "End Location": proyecto.ubicacion_final,
                "Customer PM": proyecto.gerente_proyecto or "",
                "Customer PO": proyecto.orden_compra  or "",
                 }

            # Función para verificar si una celda está en un rango combinado y obtener la celda superior izquierda
            def get_top_left_cell(cell_coordinate):
                for merged_range in ws.merged_cells.ranges:
                    min_col, min_row, max_col, max_row = range_boundaries(str(merged_range))
                    if ws[cell_coordinate].row >= min_row and ws[cell_coordinate].row <= max_row and ws[cell_coordinate].column >= min_col and ws[cell_coordinate].column <= max_col:
                        top_left_cell = ws.cell(row=min_row, column=min_col)
                        return top_left_cell
                return ws[cell_coordinate]

            # Función para asignar valor a la celda, manejando celdas combinadas
            def assign_value(cell_coordinate, value):
                top_left_cell = get_top_left_cell(cell_coordinate)
                top_left_cell.value = value

            # Registrar los datos específicos en las filas 1 y 2
            assign_value('D1', specific_data["Project"])
            assign_value('H1', specific_data["Start Date"])
            assign_value('K1', specific_data["End Date"])
            assign_value('N1', specific_data["Customer PM"])

            assign_value('D2', specific_data["Vessel"])
            assign_value('H2', specific_data["Start Location"])
            assign_value('K2', specific_data["End Location"])
            assign_value('N2', specific_data["Customer PO"])

            # Obtener los datos dinámicos del modelo DetalleProyecto
            detalles_proyectos = DetalleProyecto.objects.all()

            # Mapeo de datos a la estructura esperada
            data = []
            for detalle in detalles_proyectos:
                data.append([
                    detalle.apellido_persona,
                    detalle.nombre_persona,
                    detalle.puesto_trabajo,
                    detalle.pais_ciudadania,
                    detalle.aeropuerto_cercano,
                    detalle.dob.strftime("%d/%m/%Y") if detalle.dob else '',
                    detalle.numero_pasaporte,
                    detalle.fecha_exp_pasaporte.strftime("%d/%m/%Y") if detalle.fecha_exp_pasaporte else '',
                    detalle.tipo_visa,
                    detalle.numero_visa,
                    detalle.fecha_exp_visa.strftime("%d/%m/%Y") if detalle.fecha_exp_visa else '',
                    detalle.fecha_inicio_persona.strftime("%d/%m/%Y") if detalle.fecha_inicio_persona else '',
                    detalle.fecha_final_persona.strftime("%d/%m/%Y") if detalle.fecha_final_persona else '',
                    "NO" if not detalle.confirmado else "YES"
                ])

            # Especificar la fila y columna iniciales para los registros
            start_row = 5
            start_col = 2  # Columna B

            # Agregar cada registro a la hoja de cálculo
            for idx, record in enumerate(data):
                for col_idx, value in enumerate(record, start=start_col):
                    ws.cell(row=start_row + idx, column=col_idx, value=value)

            ws.column_dimensions[get_column_letter(6)].width = 40  # Cambia 35 por el ancho deseado

            # Ajustar la columna C para que se vea bien el nombre
            for row in range(start_row, start_row + len(data)):
                cell = ws.cell(row=row, column=3)
                cell.alignment = Alignment(vertical='center', wrap_text=True)

            # Guardar el archivo Excel en el ZIP
            excel_file = BytesIO()
            wb.save(excel_file)
            excel_file.seek(0)
            zip_file.writestr(f"{folder_general}/USS_Technicians_Spreadsheet.xlsx", excel_file.read())
            
            
            # --- Documentos Word Cabin request ---
            # Obtener los registros de DetalleProyecto

            detalles = DetalleProyecto.objects.all()
            persona = Persona.objects.all()

            # Verifica si hay detalles disponibles
            if not detalles.exists():
                return HttpResponse("No hay detalles disponibles para este proyecto.")

            # Crear un objeto en memoria para el archivo
            buffer = io.BytesIO()

            # Crear un nuevo documento
            doc = Document()

            # Configurar la página en orientación horizontal
            section = doc.sections[0]
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_height, section.page_width = section.page_width, section.page_height

            # Ajustar márgenes
            section.top_margin = Pt(10)
            section.bottom_margin = Pt(10)
            section.left_margin = Pt(35)
            section.right_margin = Pt(35)

            # Añadir encabezados centralizados con formato específico
            self.add_centered_heading(doc, 'CABIN Request', 13.5)
            self.add_centered_heading(doc, 'Vendors / Contractors', 9.5)
            self.add_centered_heading(doc, '(Information requested is important. Please fill out completely. Thank you.)', 10.5)

            # Crear tabla principal con ancho ajustado para una página
            table = doc.add_table(rows=1, cols=10)
            table.autofit = False  # Desactivar el ajuste automático de tamaño de la tabla

            # Ajustar la anchura de las columnas para que ocupen todo el ancho de la página
            widths = [Inches(1.0) for _ in range(10)]
            for i, width in enumerate(widths):
                for cell in table.columns[i].cells:
                    cell.width = width

            # Establecer los encabezados
            hdr_cells = table.rows[0].cells
            headers = [
                'Date of EMBARKATION\nmm/dd/yyyy',
                'Date of DEBARKATION\nmm/dd/yyyy',
                'COMPANY',
                'LAST NAME',
                'FIRST NAME',
                'NATIONALITY',
                'Passport Number/Country of Issue',
                'D.O.B.\nmm/dd/yyyy',
                'Contractors Email & Brief Description of work to be done',
                'Brief Description of work to be done.'
            ]

            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                self.set_cell_border(hdr_cells[i])
                self.set_font(hdr_cells[i], bold=True)

            # Añadir los datos del modelo en una sola fila por registro
            for detalle in detalles:
                #Validar si existe una relacion en los modelos DetalleProyecto y Personal
                persona = Persona.objects.get(codigo_persona=detalle.codigo_persona)
                row_cells = table.add_row().cells
                row_data = [
                    detalle.fecha_inicio_persona.strftime('%m/%d/%Y') if detalle.fecha_inicio_persona else '',
                    detalle.fecha_final_persona.strftime('%m/%d/%Y') if detalle.fecha_final_persona else '',
                    detalle.companhia or "", 
                    detalle.apellido_persona or '',
                    detalle.nombre_persona or '',
                    detalle.pais_ciudadania or '',
                    f"{detalle.numero_pasaporte}/{persona.id_pais_emision_doc.nombre_pais}" or '',
                    detalle.dob.strftime('%m/%d/%Y') if detalle.dob else '',
                    detalle.email_persona or '',  
                    detalle.puesto_trabajo or ''
                ]

                for i, value in enumerate(row_data):
                    row_cells[i].text = value
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    self.set_cell_border(row_cells[i])
                    self.set_font(row_cells[i])


            # Aplicar bordes alrededor de la tabla (cuadro)
            #self.set_table_border(table)

            # Añadir un salto de línea después de la tabla
            doc.add_paragraph()

            # Añadir la línea "Date of Request" y "SHIP NAME"
            paragraph = doc.add_paragraph()
            run = paragraph.add_run('Date of Request: ')
            run.font.size = Pt(14)
            run.bold = True
            run = paragraph.add_run(f'__{ datetime.now().strftime("%m/%d/%Y")}_____')
            run.font.size = Pt(10)
            run.bold = True
            run = paragraph.add_run('\tSHIP NAME: ')
            run.font.size = Pt(13)
            run.bold = True
            run = paragraph.add_run('_______CELEBRITY INFINITY_________')
            run.font.size = Pt(10)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Añadir la línea "Additional Comments and/or Instructions"
            paragraph = doc.add_paragraph()
            run = paragraph.add_run('Additional Comments and/or Instructions:')
            run.font.size = Pt(9.5)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Guardar el archivo Word en el ZIP
            word_cabin_request = BytesIO()
            doc.save(word_cabin_request)
            word_cabin_request.seek(0)
            zip_file.writestr(f"{folder_general}/Cabin_request_.docx", word_cabin_request.read())
    
            
            # --- Documentos Word cleance_Letter ---
            
            for detalle in detalles_proyectos:
                doc = Document()

                # Encabezado
                section = doc.sections[0]
                header = section.header
                header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
                header_table.autofit = False
                header_table.columns[0].width = Inches(2.5)
                header_table.columns[1].width = Inches(4.0)
                hdr_cells = header_table.rows[0].cells
                run = hdr_cells[0].paragraphs[0].add_run()
                run.add_picture('./static/img/logo.png', width=Inches(2.0))

                # Fecha actual
                current_date = datetime.now().strftime("%B %d, %Y")
                paragraph = hdr_cells[0].paragraphs[0]
                run = paragraph.add_run(f"\n{current_date}")
                run.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)

                # Añadir dirección
                hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                paragraph = hdr_cells[1].paragraphs[0]
                run = paragraph.add_run("\n\nUNITED SHIP SERVICE INC.\n")
                run.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)

                run = paragraph.add_run("PALM GROVE HOUSE P.O. BOX 438\n")
                run.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)

                run = paragraph.add_run("Road Town, Tortola British Virgin Islands")
                run.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)

                # Añadir un párrafo para la dirección del inspector
                paragraph = doc.add_paragraph("\n\n")
                run = paragraph.add_run("Immigration Inspector")
                run.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(11)
                paragraph.add_run("\nRe: Clearance Letter")

                # Añadir el saludo
                paragraph = doc.add_paragraph("\n\nDear Sir or Madam,\n")
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.name = 'Calibri Light'
                paragraph.runs[0].font.size = Pt(11)

                # Añadir el cuerpo principal de la carta
                body_text = (
                    "This letter will serve to confirm that the person detailed below is a "
                    "representative of "
                )
                paragraph = doc.add_paragraph(body_text)
                paragraph.runs[0].font.name = 'Calibri'
                paragraph.runs[0].font.size = Pt(11)
                run = paragraph.add_run("UNITED SHIP SERVICE INC.")
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run = paragraph.add_run(" and has been requested to board the vessel ")
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run = paragraph.add_run("CELEBRITY INFINITY on JULY 20 – AUGUST 03, 2024, at Port of Piraeus, Greece.")
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                paragraph.add_run("\n\nPlease allow the Representative to enter with his ESTA or VISA.\n")

                # Detalles del representante
                table = doc.add_table(rows=2, cols=5)
                table.style = 'Table Grid'
                
                # Desactivar auto ajuste
                table.autofit = False

                # Ajustar el ancho de las columnas manualmente
                table.columns[0].width = Inches(2.0)  # LAST NAME
                table.columns[1].width = Inches(2.0)  # FIRST NAME
                table.columns[2].width = Inches(1.5)  # DOB
                table.columns[3].width = Inches(1.5)  # NATIONALITY
                table.columns[4].width = Inches(2.0)  # PASSPORT


                # Aplicar tamaño de fuente 10 a todas las celdas de la tabla
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                headers = ['LAST NAME', 'FIRST NAME', 'DOB', 'NATIONALITY', 'PASSPORT']
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                    # Color de fondo
                    tc = hdr_cells[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '4F4F4F')
                    tcPr.append(shd)

                # Añadir datos del detalle
                row_cells = table.rows[1].cells
                row_cells[0].text = detalle.apellido_persona or ""
                row_cells[1].text = detalle.nombre_persona or ""
                row_cells[2].text = detalle.dob.strftime("%m/%d/%Y") if detalle.dob else ""
                row_cells[3].text = detalle.pais_ciudadania or ""
                row_cells[4].text = detalle.numero_pasaporte or ""

                # Cambiar el color del texto y poner en negrita
                hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)


                # Añadir el siguiente párrafo
                paragraph = doc.add_paragraph()
                paragraph.add_run(
                    "\nAs all operational Cruise vessels have 'Ship Systems' predominantly of European "
                    "manufacture it is therefore a requirement to regularly update the systems to meet with "
                    "USCG and IMO-MARPOL certifications. Due to the nature of the upgrading our personnel must "
                    "work on the vessel during normal operation within International Waters and/or dry-dock facilities "
                    "outside the Continental USA.\n\n"
                )
                paragraph.add_run(
                    "When such service is required we will cover all costs and bear full responsibility for the "
                    "above-mentioned person during his stay onboard.\n\n"
                )
                paragraph.add_run(
                    "Please do not hesitate to contact the undersigned should you require further information and/or "
                    "clarification of this request."
                )
                paragraph.style.font.name = 'Calibri'
                paragraph.style.font.size = Pt(10)

                # Añadir firma, nombre y título en una tabla
                signature_table = doc.add_table(rows=2, cols=1)
                signature_cells = signature_table.rows[0].cells

                # Colocar la imagen de la firma en la primera fila
                run = signature_cells[0].paragraphs[0].add_run()
                run.add_picture('./static/img/firma.png', width=Inches(2.0))

                # Colocar el nombre, título y contacto en la segunda fila
                signature_cells = signature_table.rows[1].cells
                paragraph = signature_cells[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = paragraph.add_run("Henry J López\n")
                run.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
                paragraph.add_run("Project Administrator\n")
                paragraph.add_run("Cell Phone: +1 561 568 0140\n")
                paragraph.add_run("Email: hlopez@uss-us.com")

                # Eliminar los bordes de la tabla de la firma
                tbl = signature_table._tbl
                tblPr = tbl.tblPr
                tblBorders = tblPr.find(qn('w:tblBorders'))

                if tblBorders is not None:
                    tblPr.remove(tblBorders)

                # Configurar el pie de página
                footer = section.footer
                footer_paragraph = footer.paragraphs[0]
                footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Añadir texto al pie de página
                footer_paragraph.clear()
                run = footer_paragraph.add_run(
                    "U.S. REPRESENTATIVE - UNITED SHIP SERVICE, CORP. 2835 N Commerce Pkwy, Miramar, FL 33025,\n"
                    "USA PHONE :+1 954 583 4588 • FAX: +1 954 583 4906 • CELL: +1 954 559 6018 • E-MAIL: info@uss-us.com"
                )
                run.font.name = 'Calibri Light'
                run.font.size = Pt(8)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri Light')

                # Ajustar el tamaño de las secciones y márgenes para que el contenido quede en una sola página
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

                # Guardar el archivo Word en el ZIP
                word_file = BytesIO()
                doc.save(word_file)
                word_file.seek(0)

                # Procesar imágenes
                # validar la primera coicidencia 
                persona = Persona.objects.filter(codigo_persona=detalle.codigo_persona).first()
            
                if persona:
                        # Guardar imagen de la persona
                        if persona.imagen_persona1:  # Asegúrate de que esta es la imagen correcta
                            imagen_path = persona.imagen_persona1.path
                            zip_file.write(imagen_path, arcname=f"{folder_general}/{detalle.nombre_persona}_{detalle.apellido_persona}/imagen_persona.jpg")

                        # Guardar imagen del pasaporte
                        if persona.imagen_pasaporte1:
                            pasaporte_path = persona.imagen_pasaporte1.path
                            zip_file.write(pasaporte_path, arcname=f"{folder_general}/{detalle.nombre_persona}_{detalle.apellido_persona}/imagen_pasaporte.jpg")

                        # Guardar imagen de la visa
                        if persona.imagen_visa1:
                            visa_path = persona.imagen_visa1.path
                            zip_file.write(visa_path, arcname=f"{folder_general}/{detalle.nombre_persona}_{detalle.apellido_persona}/imagen_visa.jpg")

            
                  
                persona = Persona.objects.filter(codigo_persona=detalle.codigo_persona).first()
                if persona:
                    # Ruta del archivo PDF original
                    input_pdf_path = os.path.join(settings.BASE_DIR, 'USS_contractor_package_2024.pdf')

                    pdf_document = fitz.open(input_pdf_path)

                    # Seleccionar la primera página del PDF
                    page = pdf_document[0]

                    # Campos y valores a insertar
                    fields_to_fill = {
                        "First Name": persona.nombre_persona or "",
                        "Last Name": persona.apellido_persona or "",
                        "Email Address": persona.email_persona1 or "",
                        "Date": persona.fecha_ingreso.strftime('%d-%m-%Y') or "",  # Asegúrate de que este campo exista
                        "Street Address": persona.direccion_persona or "",
                        "Home Phone Number": persona.telefono_persona or "",
                        "Mobile Phone Number": persona.telefmov_persona1 or "",  # Asegúrate de que este campo exista
                        "City": persona.ciudad_residencia or "",  # Asegúrate de que este campo exista
                        "Country": detalle.pais_ciudadania or "",  # Asegúrate de que este campo exista
                        "Postal Code": persona.codigo_postal or "",  
                        "Home Airport": persona.aeropuerto_cercano or "",
                        "Date of Birth": persona.fecha_nacimiento.strftime('%d-%m-%Y') or "",  
                        "Country of Birth":persona.id_pais_nacimiento.nombre_pais or "",  
                        "Country of Citizenship": persona.id_pais_nacionalidad.nombre_pais or "",  
                        "National Id #": persona.nro_doc_identidad or "",  # trabaja con el numero dde pasaporte
                        
                        #Personal Appearance & Work Clothes Sizes
                        "Hair Color": "Hair Color" or "",
                        "Eye Color": "Eye Color" or "",
                        "Height": f"{persona.estatura} Ft" or "",
                        "Weight": f"{persona.peso} Lb" or "",
                        "T-shirt Size": persona.talla_tshirt or "",
                        "Coverall Size": persona.talla_coverall or "",
                        "Pant Size": persona.talla_pantalon or "",

                        #Emergency Contact Information
                        "First & Last Name": persona.nombre_apellidos_contacto or "",
                        "Phone Number": persona.telefono_contacto or "",
                        "Email Contact": persona.email_contacto or "",

                        #Passport, Visa and Travel Documents
                        "Passport Issued by Country": persona.id_pais_emision_doc.nombre_pais or "" ,
                        "Passport Number": detalle.numero_pasaporte or "",
                        "Passport Issue Date": persona.fecha_emision_doc or "",
                        "Passport Expiration Date": persona.fecha_vencimiento_doc or "",
                        
                        #Passport Issued
                        "C1D": detalle.tipo_visa or "",
                        "B1B2": detalle.tipo_visa or "",
                        "H1B or L1": detalle.tipo_visa or "",
                        "ARC": detalle.tipo_visa or "",
                        "C1D Exp Date": detalle.fecha_exp_visa.strftime('%d-%m-%Y-') or "",
                        "B1B2 Exp Date": detalle.fecha_exp_visa.strftime('%d-%m-%Y') or "",
                        "H1B or L1 Exp Date": detalle.fecha_exp_visa.strftime('%d-%m-%Y') or "",
                        "ARC Exp Date": detalle.fecha_exp_visa.strftime('%d-%m-%Y') or "",

                        #Education & Work Experience
                        "Profession": persona.id_especialidad1.nombre_especialidad or "",
                        "Other Skills": "Other Skills" or "",
                        "Certificates": "Certificates" or "",
                        "Education/ Diplomas": "Education/ Diplomas" or "",

                        #Bank Information
                        "Bank Name": persona.id_banco.nombre_banco if persona.id_banco else "",#
                        "Bank Address": persona.direccion_banco or "", 
                        "City Bank":persona.ciudad_banco  or "",
                        "Country Bank": persona.id_pais_banco.nombre_pais if persona.id_banco else "",
                        "Postal Cod": persona.codigo_postal_banco or "",
                        "Telephone Number": persona.telefmov_persona1 or "",
                        "Bank Account Number": persona.cuenta_bancaria or "" ,
                        "Bank IBAN Number": "Bank IBAN Number" or "",
                        "Fax Number": "Fax Number" or "",
                        "Bank SWIFT Code or ABA Code": persona.codigo_swift ,
                        "Bank SORT Code": persona.codigo_corto_banco or "", 
                        "Bank Clearing Number": "Bank Clearing Number" or "",
                        "Name of the Person on the Bank Account": persona.titular_cuenta_bancaria or "",
                        "Address for the Person on the Bank Account": persona.direccion_banco or "",
                    }

                    # Insertar los valores en las posiciones específicas de la página
                    # Personal Information & Home address
                    # fila 1
                    page.insert_text((42, 144), fields_to_fill['First Name'], fontsize=10)
                    page.insert_text((175, 144), fields_to_fill['Last Name'], fontsize=10)
                    page.insert_text((308, 144), f"{fields_to_fill['Email Address']}", fontsize=9)
                    page.insert_text((475, 144), fields_to_fill['Date'], fontsize=10)

                    # fila 2
                    page.insert_text((42, 172), fields_to_fill['Street Address'], fontsize=8)
                    page.insert_text((308, 172), fields_to_fill['Home Phone Number'], fontsize=10)
                    page.insert_text((440, 172), fields_to_fill['Mobile Phone Number'], fontsize=10)

                    # fila 3
                    page.insert_text((42, 198), fields_to_fill['City'], fontsize=10)
                    page.insert_text((175, 198), fields_to_fill['Country'], fontsize=10)
                    page.insert_text((308, 198), fields_to_fill['Postal Code'], fontsize=10)
                    page.insert_text((440, 198), f"{fields_to_fill['Home Airport']}", fontsize=8)

                    # fila 4
                    page.insert_text((42, 226), fields_to_fill['Date of Birth'], fontsize=10)
                    page.insert_text((175, 226), fields_to_fill['Country of Birth'], fontsize=10)
                    page.insert_text((308, 226), fields_to_fill['Country of Citizenship'], fontsize=10)
                    page.insert_text((440, 226), fields_to_fill['National Id #'], fontsize=10)

                    #Personal Appearance & Work Clothes Sizes
                    #fila 1
                    page.insert_text((42, 275), f"{fields_to_fill['Hair Color']}", fontsize=10)
                    page.insert_text((175, 275), f"{fields_to_fill['Eye Color']}", fontsize=10)
                    page.insert_text((308, 275), f"{fields_to_fill['Height']}", fontsize=10)
                    page.insert_text((440, 275), f"{fields_to_fill['Weight']}", fontsize=10)
                    
                    # Coordenadas para las casillas de verificación y campos de texto
                    tshirt_sizes = {    "S": (61, 306), "M": (109, 306), "L": (151, 306), 
                                        "XL": (61, 328), "2XL": (109, 328), "3XL": (151, 328)}

                    coverall_sizes = {  "48": (199, 306), "50": (241, 306), "52": (280, 306),
                                        "54": (199, 328),"56": (241, 328),"58": (281, 328),}
                    
                    pant_sizes = {  "30:32": (332, 306), "32:32": (368, 306), "34:32": (405, 306),"34:34": (441, 307),
                                    "34:36": (477, 306), "36:32": (513, 306), "36:34": (549, 307),"36:36": (332, 328), 
                                    "38:32": (368, 328), "38:34": (405, 328),"38:36": (441, 328),"40:32": (477, 328), 
                                    "40:34": (513, 328), "40:36": (549, 328),}
                    
                    # Insertar tamaño de T-shirt
                    page.insert_text(tshirt_sizes[fields_to_fill["T-shirt Size"]], "x", fontsize=20, overlay=True)
                    # Insertar tamaño de Coverall
                    page.insert_text(coverall_sizes[fields_to_fill["Coverall Size"]], "x", fontsize=20, overlay=True)
                    # Insertar tamaño de Pant
                    page.insert_text(pant_sizes[fields_to_fill["Pant Size"]], "x", fontsize=20, overlay=True)

                                                                                              
                    #Emergency Contact Information
                    #fila 1
                    page.insert_text((42, 380), f"{fields_to_fill['First & Last Name']}", fontsize=10)
                    page.insert_text((271, 380), f"{fields_to_fill['Phone Number']}", fontsize=10)
                    page.insert_text((390, 380), f"{fields_to_fill['Email Contact']}", fontsize=10)

                    #Passport, Visa and Travel Documents
                    #fila 1
                    page.insert_text((42, 430), f"{fields_to_fill['Passport Issued by Country']}", fontsize=10)
                    page.insert_text((175, 430), f"{fields_to_fill['Passport Number']}", fontsize=10)
                    page.insert_text((308, 430), f"{fields_to_fill['Passport Issue Date']}", fontsize=10)
                    page.insert_text((440, 430), f"{fields_to_fill['Passport Expiration Date']}", fontsize=10)

                    #fila 2
                    visa_coordinates = {
                            "C1D Visa Yes": (124, 448), "C1D Visa No": (155, 448),
                            "B1B2 Visa Yes": (259, 448), "B1B2 Visa No": (290, 448),
                            "H1B or L1 Yes": (389, 448), "H1B or L1 No": (424, 448),
                            "ARC Yes": (523, 448), "ARC No": (556, 448)
                        }
                    # Insertar Visa C1D y su fecha de expiración
                    if fields_to_fill["C1D"] == "C1D":
                        page.insert_text(visa_coordinates["C1D Visa Yes"], "x", fontsize=20)
                        page.insert_text((82, 475), fields_to_fill["C1D Exp Date"], fontsize=10)
                    else:
                        page.insert_text(visa_coordinates["C1D Visa No"], "x", fontsize=20)

                    # Insertar Visa B1 o B1B2
                    if fields_to_fill["B1B2"] == "B1B2" or fields_to_fill["B1B2"] == "B1" :
                        page.insert_text(visa_coordinates["B1B2 Visa Yes"], "x", fontsize=20)
                        page.insert_text((218, 475), fields_to_fill["B1B2 Exp Date"], fontsize=10)
                    else:
                        page.insert_text(visa_coordinates["B1B2 Visa No"], "x", fontsize=20)
                    
                    # Insertar Visa H1B o L1
                    if fields_to_fill["H1B or L1"] == "H1B" or fields_to_fill["H1B or L1"] == "L1":
                        page.insert_text(visa_coordinates["H1B or L1 Yes"], "x", fontsize=20)
                        page.insert_text((348, 475), fields_to_fill["H1B or L1 Exp Date"], fontsize=10)
                    else:
                        page.insert_text(visa_coordinates["H1B or L1 No"], "x", fontsize=20)
                    
                    # Insertar ARC y su fecha de expiración
                    if fields_to_fill["ARC"] == "ARC":
                        page.insert_text(visa_coordinates["ARC Yes"], "x", fontsize=20)
                        page.insert_text((483, 475), fields_to_fill["ARC Exp Date"], fontsize=10)
                    else:
                        page.insert_text(visa_coordinates["ARC No"], "x", fontsize=20)

                    #Education & Work Experience
                    # fila 1
                    page.insert_text((42, 530), f"{fields_to_fill['Profession']}", fontsize=10)
                    page.insert_text((308, 530), f"{fields_to_fill['Other Skills']}", fontsize=10)
                    page.insert_text((42, 558), f"{fields_to_fill['Certificates']}", fontsize=10)
                    page.insert_text((308, 558), f"{fields_to_fill['Education/ Diplomas']}", fontsize=10)

                    #Bank Information
                    #fila 1
                    page.insert_text((42, 608), f"{fields_to_fill['Bank Name']}", fontsize=10)
                    page.insert_text((308, 608), f"{fields_to_fill['Bank Address']}", fontsize=8)

                    #fila 2
                    page.insert_text((42, 635), f"{fields_to_fill['City Bank']}", fontsize=10)
                    page.insert_text((175, 635), f"{fields_to_fill['Country Bank']}", fontsize=10)
                    page.insert_text((308, 635), f"{fields_to_fill['Postal Cod']}", fontsize=10)
                    page.insert_text((440, 635), f"{fields_to_fill['Telephone Number']}", fontsize=10)

                    #fila 3
                    page.insert_text((42, 660), f"{fields_to_fill['Bank Account Number']}", fontsize=10)
                    page.insert_text((308, 660), f"{fields_to_fill['Bank IBAN Number']}", fontsize=10)
                    page.insert_text((440, 660), f"{fields_to_fill['Fax Number']}", fontsize=10)

                    #fila 4
                    page.insert_text((42, 690), f"{fields_to_fill['Bank SWIFT Code or ABA Code']}", fontsize=10)
                    page.insert_text((308, 690), f"{fields_to_fill['Bank SORT Code']}", fontsize=10)
                    page.insert_text((440, 690), f"{fields_to_fill['Bank Clearing Number']}", fontsize=10)

                    #fila5
                    page.insert_text((42, 715), f"{fields_to_fill['Name of the Person on the Bank Account']}", fontsize=10)
                    page.insert_text((308, 715), f"{fields_to_fill['Address for the Person on the Bank Account']}", fontsize=8)

                    # Guardar el PDF en el zip
                    pdf_buffer = BytesIO()
                    pdf_document.save(pdf_buffer)
                    pdf_filename = f"USS_contractor_package_{datetime.now().year}.pdf"
                    zip_file.writestr(f"{folder_general}/{detalle.nombre_persona}_{detalle.apellido_persona}/{pdf_filename}", pdf_buffer.getvalue())
                    pdf_document.close()


                folder_name = f"{detalle.nombre_persona}_{detalle.apellido_persona}/"
                doc_filepath = f"{folder_general}/{folder_name}clearance_letter_{detalle.nombre_persona}_{detalle.apellido_persona}.docx"
                zip_file.writestr(doc_filepath, word_file.read())

        zip_buffer.seek(0)

        # Configurar la respuesta HTTP para descargar el archivo ZIP
        response = HttpResponse(zip_buffer, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{proyecto}.zip"'

        return response